from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
    hs.font.name = 'Calibri'

# Helper functions
def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_note(text, label="NOTE"):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    p.add_run(text)

# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Asana")
run.font.size = Pt(42)
run.bold = True
run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Yoga Pose Chatbot")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)

doc.add_paragraph()

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Product Requirements Document")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version 1.0  |  May 2026  |  Status: Draft").font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'Asana is an AI-powered yoga pose chatbot that acts as a personal, always-available yoga instructor. '
    'It guides practitioners — from absolute beginners to intermediates — through yoga poses with step-by-step '
    'instructions, safety cues, modifications and benefits. It currently exists as a working prototype and needs '
    'to be re-built from scratch for production.'
)
doc.add_paragraph(
    'This PRD defines the what and why; the tech team decides the how (languages, frameworks, libraries). '
    'A recommended system design is included to set architectural guardrails.'
)

# ===== 2. USERS & PERSONAS =====
doc.add_heading('2. Users & Personas', level=1)
add_table(
    ['Persona', 'Description', 'Needs'],
    [
        ['Beginner Yogi (Priya)', 'New to yoga, practises at home without a teacher. Might ask in Hindi/Hinglish.',
         'Safe, simple instructions; modifications; injury warnings; voice input for hands-free use'],
        ['Intermediate Practitioner (Alex)', '1-2 years of yoga, exploring new poses. English speaker.',
         'Deeper alignment cues, sequencing advice, Sanskrit names, detailed benefits'],
        ['Yoga Teacher (Meera)', 'Uses the app as a quick reference while planning classes.',
         'Teaching tips, cueing language, common student mistakes, sequencing guidance'],
    ]
)

doc.add_heading('Out of Scope Users (v1)', level=2)
add_bullet('Advanced/professional yogis requiring extremely nuanced adjustments')
add_bullet('Users seeking nutrition, meditation, or general fitness advice (Asana will explicitly redirect these)')

# ===== 3. PROBLEM STATEMENT =====
doc.add_heading('3. Problem Statement', level=1)
doc.add_paragraph(
    'Millions of people practise yoga at home without professional guidance. The risks are:'
)
add_bullet('Incorrect alignment → injuries (knee, back, wrist, neck)', 'Incorrect alignment: ')
add_bullet('No personalised modifications → frustration and dropout for those with physical limitations', 'No personalised modifications: ')
add_bullet('Information overload → YouTube/Instagram offer contradictory advice with no way to ask follow-up questions', 'Information overload: ')
add_bullet('Language barrier → Many Indian users think in Hindi/Hinglish but most quality yoga content is English-only', 'Language barrier: ')
doc.add_paragraph()
doc.add_paragraph(
    'Existing solutions (generic chatbots, search engines, YouTube) are either too shallow, non-interactive, or not safety-focused.'
)

# ===== 4. SOLUTION =====
doc.add_heading('4. Solution', level=1)
doc.add_paragraph('Asana is a conversational AI yoga pose guide that:')
add_bullet('Answers pose-specific questions using a curated, expert-vetted knowledge base (not hallucinated internet content)')
add_bullet('Uses Retrieval-Augmented Generation (RAG) to ground every response in verified yoga data')
add_bullet('Provides a structured, consistent response format (name, steps, benefits, mistakes, modifications, duration)')
add_bullet('Supports voice input and text-to-speech for hands-free use during practice')
add_bullet('Understands Hindi/Hinglish input and responds in English')
add_bullet('Maintains conversation memory for follow-up questions within a session')
add_bullet('Refuses out-of-scope questions gracefully (nutrition, meditation, general fitness)')

# ===== 5. FEATURES & PRIORITISATION =====
doc.add_heading('5. Features & Prioritisation (MoSCoW)', level=1)

doc.add_heading('Must Have (P0) — Launch Blockers', level=2)
add_table(
    ['#', 'Feature', 'Description'],
    [
        ['F1', 'Pose Q&A with RAG', 'User asks about any yoga pose → structured response with instructions, benefits, mistakes, modifications, duration. Grounded in curated knowledge base.'],
        ['F2', 'Curated Knowledge Base', 'Expert-vetted yoga data covering 20+ poses across 4 categories: standing, seated, floor/supine, inversions & balance. Includes teaching tips.'],
        ['F3', 'Structured Response Format', 'Every pose response follows: Pose Name (English/Sanskrit) → Instructions → Benefits → Common Mistakes → Modifications → Duration'],
        ['F4', 'Conversation Memory', 'Session-based chat history so users can ask follow-ups'],
        ['F5', 'Scope Guardrails', 'Bot politely declines non-yoga-pose questions and redirects. Admits when it doesn\'t have data on a pose.'],
        ['F6', 'Safety-First Responses', 'All responses include contraindication warnings where applicable (e.g., headstand → no if high BP, neck injury, pregnancy)'],
        ['F7', 'Text Chat Interface', 'Clean, intuitive chat UI for text-based conversations'],
        ['F8', 'API Backend', 'RESTful API serving chat requests with session management'],
    ]
)

doc.add_heading('Should Have (P1) — Target for v1.0', level=2)
add_table(
    ['#', 'Feature', 'Description'],
    [
        ['F9', 'Voice Input (STT)', 'Hold-to-speak microphone for hands-free pose queries. Supports Hindi and English.'],
        ['F10', 'Voice Output (TTS)', 'Auto-speak bot responses so user doesn\'t have to look at screen mid-pose'],
        ['F11', 'Hindi/Hinglish Understanding', 'Accept queries in Hindi or Hinglish, respond in English'],
        ['F12', 'Mobile App (Android)', 'Native Android app with splash screen, chat screen, and voice recording screen'],
        ['F13', 'Typing Indicator', 'Animated indicator while AI is generating a response'],
    ]
)

doc.add_heading('Could Have (P2) — Post-Launch Enhancements', level=2)
add_table(
    ['#', 'Feature', 'Description'],
    [
        ['F14', 'iOS App', 'Native iOS client'],
        ['F15', 'Pose Image/Video Attachments', 'Show illustrative images or short video clips alongside pose instructions'],
        ['F16', 'Personalised Sequences', 'Generate a sequenced practice based on the knowledge base'],
        ['F17', 'User Profiles & History', 'Persistent user accounts, saved favourite poses, practice history'],
        ['F18', 'Multi-language Responses', 'Respond in Hindi, Tamil, etc.'],
        ['F19', 'Web Client', 'Browser-based chat interface'],
        ['F20', 'Analytics Dashboard', 'Track most-asked poses, user engagement, error rates'],
    ]
)

doc.add_heading("Won't Have (v1)", level=2)
add_table(
    ['#', 'Feature', 'Reason'],
    [
        ['W1', 'Pose detection via camera', 'Requires computer vision; different product vertical'],
        ['W2', 'Live class integration', 'Requires video streaming infrastructure'],
        ['W3', 'E-commerce / class booking', 'Out of scope for an AI chatbot'],
        ['W4', 'Social features', 'Out of scope for v1'],
    ]
)

# ===== 6. SYSTEM DESIGN =====
doc.add_heading('6. System Design', level=1)

doc.add_heading('6.1 High-Level Architecture', level=2)
doc.add_paragraph(
    'The system follows a standard three-tier architecture: Client Layer → API Gateway → Application Layer → Data Layer, '
    'with an offline data pipeline for knowledge base ingestion.'
)

doc.add_heading('Architecture Components', level=3)
add_table(
    ['Layer', 'Components', 'Purpose'],
    [
        ['Client Layer', 'Android App (Chat + Voice), iOS App (Future), Web Client (Future)', 'User-facing interfaces'],
        ['API Gateway', 'Load Balancer (NGINX / Cloud LB)', 'SSL termination, rate limiting, routing, health checks'],
        ['Application Layer', 'API Server (REST) + RAG Engine + Session Manager', 'Core business logic'],
        ['AI / ML Layer', 'LLM Provider (Gemini/GPT/Claude) + Embedding Model', 'Intelligence layer'],
        ['Data Layer', 'Vector DB + Knowledge Base + Cache (Redis) + Session Store', 'Persistence and retrieval'],
        ['Offline Pipeline', 'Ingestion Service + Text Splitter + Embedding Generator', 'Knowledge base processing'],
    ]
)

doc.add_heading('6.2 API Server (REST)', level=2)
doc.add_heading('Endpoints', level=3)
add_table(
    ['Method', 'Path', 'Description'],
    [
        ['GET', '/', 'Health check — returns {"status": "healthy"}'],
        ['POST', '/chat', 'Send a message, receive AI response'],
        ['DELETE', '/chat/{session_id}', 'Clear session history'],
    ]
)

doc.add_heading('POST /chat — Request Schema', level=3)
doc.add_paragraph('{\n  "session_id": "uuid-v4",\n  "message": "How do I do Warrior I pose?"\n}', style='Normal')

doc.add_heading('POST /chat — Response Schema', level=3)
doc.add_paragraph('{\n  "reply": "## Virabhadrasana I (Warrior I)\\n\\nInstructions:...",\n  "session_id": "uuid-v4",\n  "sources": ["standing_poses.txt:chunk_3"],\n  "latency_ms": 1200\n}', style='Normal')

add_note('The current prototype stores sessions in-memory (dict). Production MUST use an external session store (Redis/DynamoDB) for horizontal scaling and persistence across restarts.', 'IMPORTANT')

doc.add_heading('6.3 RAG Engine', level=2)
doc.add_paragraph('This is the core intelligence layer. For every user query:')
p = doc.add_paragraph()
p.add_run('1. ').bold = True
p.add_run('Fetch chat history for the session from Session Manager\n')
run2 = p.add_run('2. ')
run2.bold = True
p.add_run('Run similarity search on the Vector DB (query embedding, k=4 chunks)\n')
run3 = p.add_run('3. ')
run3.bold = True
p.add_run('Send System Prompt + Retrieved Context + Chat History + User Query to LLM\n')
run4 = p.add_run('4. ')
run4.bold = True
p.add_run('Return structured pose response to user\n')
run5 = p.add_run('5. ')
run5.bold = True
p.add_run('Append user message + AI response to session history')

doc.add_heading('RAG Configuration', level=3)
add_table(
    ['Parameter', 'Current Value', 'Production Recommendation'],
    [
        ['Embedding Model', 'all-MiniLM-L6-v2', 'Keep or upgrade to all-mpnet-base-v2 for better accuracy'],
        ['Chunk Size', '500 chars', '500-800 chars (one pose per chunk ideally)'],
        ['Chunk Overlap', '50 chars', '50-100 chars'],
        ['Retrieval Top-K', '4', '3-5 (test and tune)'],
        ['LLM Temperature', '0.3', '0.2-0.4 (low for accuracy, enough for natural language)'],
        ['LLM Model', 'Gemini 2.5 Flash Lite', "Any capable model — team's choice"],
    ]
)

doc.add_heading('6.4 System Prompt (Persona: "Asana")', level=2)
doc.add_paragraph('The system prompt uses the CRISPE framework:')
add_table(
    ['Element', 'Definition'],
    [
        ['Capacity', 'Certified yoga teacher with expertise in anatomy, alignment, modifications'],
        ['Role', 'Speaking with beginners to intermediate students'],
        ['Insight', 'Knowledge of 20+ poses across 4 categories'],
        ['Statement', 'Structured response format (name, steps, benefits, mistakes, modifications, duration)'],
        ['Personality', 'Professional, warm, safety-first, uses Sanskrit + English names'],
        ['Experiment', 'Consistent response template'],
    ]
)
add_note('The system prompt is a critical product asset. Any changes must be reviewed by both product and a yoga SME. Store as versioned configuration, not hardcoded.', 'IMPORTANT')

doc.add_heading('6.5 Session Manager', level=2)
add_bullet('Stores conversation history per session_id')
add_bullet('Enables multi-turn follow-up questions')
add_bullet('TTL: 30 minutes of inactivity → auto-expire')
add_bullet('Max history: 20 message pairs (40 messages) per session to manage context window')

doc.add_heading('6.6 Cache Layer', level=2)
add_bullet('Cache frequent pose queries at the response level')
add_bullet('Cache TTL: 1 hour')
add_bullet('Cache key: normalised query text (lowercased, trimmed)')
add_bullet('Invalidate on knowledge base updates')

doc.add_heading('6.7 Data Pipeline (Offline)', level=2)
doc.add_paragraph('Curated Yoga Text Files → Document Loader → Text Splitter (500 chars, 50 overlap) → Embedding Model → Vector Database')
doc.add_paragraph()

doc.add_heading('Current Knowledge Base', level=3)
add_table(
    ['File', 'Category', 'Poses Covered'],
    [
        ['standing_poses.txt', 'Standing', 'Tadasana, Warrior I, Warrior II, Triangle, Tree'],
        ['seated_poses.txt', 'Seated', 'Staff, Forward Bend, Butterfly, Seated Twist, Easy Pose'],
        ['floor_poses.txt', 'Floor / Supine', "Child's Pose, Cobra, Bridge, Corpse, Supine Twist"],
        ['inversions_and_balance.txt', 'Inversions & Balance', "Downward Dog, Headstand, Shoulder Stand, Crow, Dancer's"],
        ['teaching_tips.txt', 'Instruction', 'Cueing, sequencing, beginner tips, common mistakes, props, injury awareness'],
    ]
)
add_note('The knowledge base is the product\'s competitive moat. All data must be reviewed by a certified yoga instructor, versioned, and augmented over time (target: 50+ poses by v1.1, 100+ by v2.0).', 'WARNING')

doc.add_heading('6.8 Voice Architecture (Android)', level=2)
doc.add_paragraph('User holds mic button → Android STT (Speech-to-Text) → Transcribed text → POST /chat → AI response → Android TTS (Text-to-Speech) → Audio playback')
doc.add_paragraph()
add_bullet('STT: On-device (Android SpeechRecognizer). Primary language: hi-IN, with en-IN and en-US as additional.')
add_bullet('TTS: On-device (Android TextToSpeech). Strips markdown before speaking.')
add_bullet('UX: Hold-to-speak pattern with animated ripple feedback during recording.')

# ===== 7. NON-FUNCTIONAL REQUIREMENTS =====
doc.add_heading('7. Non-Functional Requirements', level=1)
add_table(
    ['Requirement', 'Target'],
    [
        ['Response Latency (P95)', '< 3 seconds end-to-end'],
        ['Availability', '99.5% uptime'],
        ['Concurrent Users', '500 simultaneous sessions at launch'],
        ['Horizontal Scaling', 'API server must be stateless (sessions externalized)'],
        ['Security', 'HTTPS everywhere; API keys in env vars, never in client code; rate limiting'],
        ['Data Privacy', 'No PII stored; chat sessions auto-expire; no conversation data used for training'],
        ['Cost', 'LLM cost < $0.002 per conversation turn'],
        ['Observability', 'Structured logging, request tracing, error rate monitoring, LLM latency tracking'],
        ['Mobile Performance', 'App cold start < 2s; smooth 60fps animations'],
    ]
)

# ===== 8. KEY RISKS =====
doc.add_heading('8. Key Risks & Mitigations', level=1)
add_table(
    ['Risk', 'Impact', 'Mitigation'],
    [
        ['LLM hallucination (invents a pose)', 'User injury', 'RAG grounding + fallback message + low temperature'],
        ['LLM provider outage', 'App unusable', 'Implement fallback provider; cache popular responses'],
        ['Incorrect yoga advice', 'User injury', 'Expert review of knowledge base; clear disclaimers'],
        ['High LLM costs at scale', 'Budget overrun', 'Use efficient models; response caching; rate limiting'],
        ['Voice transcription errors', 'Bad queries', 'Show transcribed text before sending; allow user to edit'],
        ['Session store failure', 'Loss of context', 'Graceful degradation; Redis cluster for HA'],
    ]
)

# ===== 9. SUCCESS METRICS =====
doc.add_heading('9. Success Metrics', level=1)
add_table(
    ['Metric', 'Target (3 months post-launch)'],
    [
        ['Daily Active Users (DAU)', '1,000+'],
        ['Avg. messages per session', '4+'],
        ['User retention (D7)', '30%+'],
        ['Response accuracy (manual audit)', '95%+ grounded in knowledge base'],
        ['NPS Score', '40+'],
        ['P95 Response Latency', '< 3s'],
        ['Crash-free sessions', '99.5%+'],
    ]
)

# ===== 10. IMPLEMENTATION ROADMAP =====
doc.add_heading('10. Suggested Implementation Roadmap', level=1)

doc.add_heading('Phase 1: Backend Foundation (Weeks 1-2)', level=2)
for item in [
    'Set up project repository and CI/CD pipeline',
    'Implement API server with /chat and health check endpoints',
    'Build RAG engine with vector DB and embedding pipeline',
    'Externalize session storage',
    'Implement system prompt (CRISPE) with versioned config',
    'Data ingestion pipeline for knowledge base',
    'Add structured logging and error handling',
    'Deploy to staging environment',
]:
    add_bullet(item)

doc.add_heading('Phase 2: Android App (Weeks 2-4)', level=2)
for item in [
    'Chat screen with message bubbles, input bar',
    'Splash screen with branding',
    'Retrofit integration with backend API',
    'Session management (generate UUID per app launch)',
    'Typing indicator animation',
    'Voice input (hold-to-speak STT)',
    'Voice output (TTS auto-read)',
    'Voice recording screen with ripple animation',
    'Error handling and offline state UI',
]:
    add_bullet(item)

doc.add_heading('Phase 3: Hardening & Launch (Weeks 4-5)', level=2)
for item in [
    'Load testing (500 concurrent sessions)',
    'Rate limiting and abuse prevention',
    'SSL/TLS and security audit',
    'Knowledge base expert review (certified yoga instructor)',
    'Add legal disclaimer ("not a substitute for professional instruction")',
    'App store listing and assets',
    'Monitoring and alerting setup',
    'Production deployment',
]:
    add_bullet(item)

doc.add_heading('Phase 4: Post-Launch (Weeks 6-8)', level=2)
for item in [
    'User feedback collection and analysis',
    'Expand knowledge base (target 50 poses)',
    'Response caching for popular queries',
    'iOS app kickoff',
    'Analytics dashboard',
]:
    add_bullet(item)

# ===== 11. OPEN QUESTIONS =====
doc.add_heading('11. Open Questions for Discussion', level=1)
questions = [
    'LLM Provider Lock-in: Should we abstract the LLM layer to support provider switching (Gemini ↔ GPT ↔ Claude) or commit to one?',
    'Knowledge Base Expansion: Who is the SME for reviewing new pose content? Do we have a certified yoga instructor on call?',
    'Disclaimer / Liability: Do we need a legal review for the "not medical advice" disclaimer? Should users accept terms before first use?',
    'Voice Quality: Should we evaluate cloud-based STT/TTS (Google Cloud STT, ElevenLabs TTS) for better quality vs on-device?',
    'Multilingual Responses: Should v1 respond in Hindi when the user writes in Hindi, or always in English?',
    'Monetisation: Is there a monetisation strategy that affects feature prioritisation (freemium, ads, subscription)?',
]
for i, q in enumerate(questions, 1):
    p = doc.add_paragraph(style='List Number')
    parts = q.split(': ', 1)
    run = p.add_run(parts[0] + ': ')
    run.bold = True
    p.add_run(parts[1])

# ===== SAVE =====
output_path = os.path.expanduser('~/Desktop/langchain/Asana_PRD.docx')
doc.save(output_path)
print(f"✅ PRD saved to: {output_path}")
